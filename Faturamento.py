import tkinter as tk
from tkinter import filedialog, messagebox
from ttkbootstrap import Style, Window, Label, Entry, Button, Frame
import pandas as pd
from datetime import datetime
from datetime import date
import os
import locale

# Certificando-se de usar a biblioteca correta
try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    messagebox.showerror("Erro de Importação", 
                         "Biblioteca python-docx não encontrada. Por favor, instale com:\n" +
                         "pip install python-docx")
    raise

# Configurar locale para formato braspileiro (R$)
try:
    locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')
except locale.Error:
    try:
        locale.setlocale(locale.LC_ALL, 'Portuguese_Brazil.1252')  # Windows
    except locale.Error:
        pass

class FaturamentoPromilApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Faturamento - Entidades PROMIL")
        self.root.geometry("600x450")
        
        # Variáveis
        self.caminho_arquivo = tk.StringVar()
        self.cnpj = tk.StringVar()
        self.data_inicial = tk.StringVar()
        self.data_final = tk.StringVar()
        self.caminho_modelo_word = tk.StringVar(value="R:\\36.FISCAL\\04.GERAL\\05.IMPOSTOS\\02.IMPOSTOS MUNICIPAIS\\01.ISSQN\\Promil Promotora (PMIL)\\Relatório de faturamento\\modelo_declaração_faturamento.docx")
        
        # Criar interface
        self.criar_interface()
    
    def criar_interface(self):
        # Frame principal
        frame_principal = Frame(self.root, padding=20)
        frame_principal.pack(fill="both", expand=True)
        
        # Título
        titulo = Label(frame_principal, text="Faturamento - Entidades PROMIL", font=("Helvetica", 16, "bold"))
        titulo.pack(pady=10)
        
        # Frame para selecionar arquivo
        frame_arquivo = Frame(frame_principal)
        frame_arquivo.pack(fill="x", pady=10)
        
        Label(frame_arquivo, text="Arquivo Excel:").pack(side="left", padx=5)
        Entry(frame_arquivo, textvariable=self.caminho_arquivo, width=40).pack(side="left", padx=5, fill="x", expand=True)
        Button(frame_arquivo, text="Selecionar", command=self.selecionar_arquivo).pack(side="left", padx=5)
        
        # Frame para modelo Word
        frame_modelo = Frame(frame_principal)
        frame_modelo.pack(fill="x", pady=10)
        
        Label(frame_modelo, text="Modelo Word:").pack(side="left", padx=5)
        Entry(frame_modelo, textvariable=self.caminho_modelo_word, width=40).pack(side="left", padx=5, fill="x", expand=True)
        Button(frame_modelo, text="Selecionar", command=self.selecionar_modelo).pack(side="left", padx=5)
        
        # Frame para CNPJ
        frame_cnpj = Frame(frame_principal)
        frame_cnpj.pack(fill="x", pady=10)
        
        Label(frame_cnpj, text="CNPJ:").pack(side="left", padx=5)
        Entry(frame_cnpj, textvariable=self.cnpj, width=20).pack(side="left", padx=5)
        
        # Frame para datas
        frame_datas = Frame(frame_principal)
        frame_datas.pack(fill="x", pady=10)
        
        Label(frame_datas, text="Data Inicial (DD/MM/AAAA):").pack(side="left", padx=5)
        Entry(frame_datas, textvariable=self.data_inicial, width=15).pack(side="left", padx=5)
        
        Label(frame_datas, text="Data Final (DD/MM/AAAA):").pack(side="left", padx=5)
        Entry(frame_datas, textvariable=self.data_final, width=15).pack(side="left", padx=5)
        
        # Botão de processamento
        Button(frame_principal, text="Processar Dados", command=self.processar_dados, bootstyle="success").pack(pady=20)
        
        # Status bar
        self.status_var = tk.StringVar()
        self.status_var.set("Pronto para processar")
        Label(self.root, textvariable=self.status_var, relief="sunken", anchor="w").pack(side="bottom", fill="x")
    
    def selecionar_arquivo(self):
        arquivo = filedialog.askopenfilename(
            title="Selecione o arquivo Excel",
            filetypes=[("Arquivos Excel", "*.xlsx;*.xls")]
        )
        if arquivo:
            self.caminho_arquivo.set(arquivo)
            self.status_var.set(f"Arquivo selecionado: {os.path.basename(arquivo)}")
    
    def selecionar_modelo(self):
        arquivo = filedialog.askopenfilename(
            title="Selecione o modelo Word",
            filetypes=[("Arquivos Word", "*.docx")]
        )
        if arquivo:
            self.caminho_modelo_word.set(arquivo)
            self.status_var.set(f"Modelo selecionado: {os.path.basename(arquivo)}")
    
    def validar_data(self, data_str):
        try:
            return datetime.strptime(data_str, "%d/%m/%Y")
        except ValueError:
            return None
    
    def formatar_data_abreviada(self, data):
        """Formata uma data no formato DD/MM/AA"""
        if isinstance(data, str):
            # Se for string, tentar converter de diferentes formatos
            try:
                data = datetime.strptime(data, "%Y-%m-%d %H:%M:%S")
            except ValueError:
                try:
                    data = datetime.strptime(data, "%Y-%m-%d")
                except ValueError:
                    try:
                        data = datetime.strptime(data, "%d/%m/%Y")
                    except ValueError:
                        return data  # Retorna a string original se não conseguir converter
        elif isinstance(data, pd.Timestamp):
            data = data.to_pydatetime()
        
        return data.strftime("%d/%m/%y")
    
    def obter_data_atual_formatada(self):
        """Obtém a data atual formatada no padrão brasileiro"""
        hoje = date.today()
        return hoje.strftime("%d/%m/%Y")
    
    def formatar_valor_brl(self, valor):
        """Formata um valor para o padrão brasileiro (R$)"""
        try:
            return locale.currency(valor, symbol=True, grouping=True)
        except:
            # Fallback se locale não funcionar
            return f"R$ {valor:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    
    def substituir_texto_word(self, doc, marcador, novo_texto):
        """Substitui um marcador por um novo texto em todo o documento Word"""
        for paragrafo in doc.paragraphs:
            if marcador in paragrafo.text:
                for run in paragrafo.runs:
                    run.text = run.text.replace(marcador, novo_texto)
        
        for tabela in doc.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in celula.paragraphs:
                        if marcador in paragrafo.text:
                            for run in paragrafo.runs:
                                run.text = run.text.replace(marcador, novo_texto)
    
    def criar_tabela_resultados(self, doc, dados):
        """Cria uma tabela com os resultados no documento Word"""
        doc.add_paragraph().add_run("Detalhamento de Faturamento").bold = True
        
        # Adicionar uma tabela com cabeçalho
        tabela = doc.add_table(rows=1, cols=3)
        
        # Definir cabeçalhos
        cabecalhos = tabela.rows[0].cells
        cabecalhos[0].text = "Data"
        cabecalhos[1].text = "HUB"
        cabecalhos[2].text = "FATURAMENTO"
        
        # Adicionar dados - apenas registros com faturamento não nulo
        for _, row in dados.iterrows():
            # Verificar se o valor de faturamento não é nulo/NaN
            if pd.notna(row['Faturamento']) and row['Faturamento'] is not None:
                linha_celulas = tabela.add_row().cells
                linha_celulas[0].text = self.formatar_data_abreviada(row['Data'])
                linha_celulas[1].text = str(row['Cidade a recolher'])
                
                # Formatar o valor para a célula
                valor_formatado = self.formatar_valor_brl(row['Faturamento']).replace('R$', '').strip()
                linha_celulas[2].text = valor_formatado
                
                # Alinhar valor à direita
                for parag in linha_celulas[2].paragraphs:
                    parag.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    def converter_para_pdf(self, caminho_word):
        """Converte o arquivo Word para PDF usando win32com ou docx2pdf"""
        caminho_pdf = caminho_word.replace('.docx', '.pdf')
        
        # Método 1: Tentar usar win32com (mais confiável no Windows)
        try:
            import win32com.client
            import pythoncom
            
            # Inicializar COM
            pythoncom.CoInitialize()
            
            word_app = None
            try:
                # Criar aplicação Word
                word_app = win32com.client.Dispatch('Word.Application')
                word_app.Visible = False
                word_app.DisplayAlerts = False
                
                # Abrir documento
                doc = word_app.Documents.Open(os.path.abspath(caminho_word))
                
                # Salvar como PDF
                doc.SaveAs2(os.path.abspath(caminho_pdf), FileFormat=17)  # 17 = PDF
                doc.Close()
                
                return caminho_pdf
                
            except Exception as e:
                raise Exception(f"Erro com win32com: {str(e)}")
            finally:
                if word_app:
                    try:
                        word_app.Quit()
                    except:
                        pass
                pythoncom.CoUninitialize()
                
        except ImportError:
            # Método 2: Tentar usar docx2pdf
            try:
                from docx2pdf import convert
                convert(caminho_word, caminho_pdf)
                return caminho_pdf
            except ImportError:
                # Método 3: Tentar usar pypandoc
                try:
                    import pypandoc
                    pypandoc.convert_file(caminho_word, 'pdf', outputfile=caminho_pdf)
                    return caminho_pdf
                except ImportError:
                    raise Exception(
                        "Nenhuma biblioteca de conversão PDF disponível.\n"
                        "Instale uma das seguintes opções:\n"
                        "- pip install pywin32 (recomendado para Windows)\n"
                        "- pip install docx2pdf\n"
                        "- pip install pypandoc"
                    )
            except Exception as e:
                raise Exception(f"Erro na conversão para PDF: {str(e)}")
    
    def processar_dados(self):
        # Verificar se todos os campos estão preenchidos
        if not self.caminho_arquivo.get():
            messagebox.showerror("Erro", "Selecione um arquivo Excel!")
            return
        
        if not self.caminho_modelo_word.get():
            messagebox.showerror("Erro", "Selecione ou informe o caminho do modelo Word!")
            return
        
        if not self.cnpj.get():
            messagebox.showerror("Erro", "Informe o CNPJ!")
            return
        
        if not self.data_inicial.get() or not self.data_final.get():
            messagebox.showerror("Erro", "Informe as datas inicial e final!")
            return
        
        # Validar datas
        data_inicial = self.validar_data(self.data_inicial.get())
        data_final = self.validar_data(self.data_final.get())
        
        if not data_inicial or not data_final:
            messagebox.showerror("Erro", "Formato de data inválido! Use DD/MM/AAAA.")
            return
        
        if data_inicial > data_final:
            messagebox.showerror("Erro", "Data inicial não pode ser maior que a data final!")
            return
        
        # Processar os dados
        try:
            self.status_var.set("Processando dados...")
            self.root.update_idletasks()
            
            # Carregar arquivo Excel
            df = pd.read_excel(self.caminho_arquivo.get(), sheet_name="ISS Promil Próprio")
            
            # Converter a coluna de data (A) para o formato datetime
            df['Data'] = pd.to_datetime(df.iloc[:, 0], errors='coerce')
            
            # Filtrar por CNPJ e data
            cnpj_filtrado = self.cnpj.get()
            dados_filtrados = df[
                (df.iloc[:, 3] == cnpj_filtrado) & 
                (df['Data'] >= data_inicial) & 
                (df['Data'] <= data_final)
            ]
            
            if dados_filtrados.empty:
                messagebox.showinfo("Informação", "Nenhum dado encontrado para os filtros informados!")
                self.status_var.set("Nenhum dado encontrado")
                return
            
            # Criar novo DataFrame para exportação
            resultado = pd.DataFrame({
                'CNPJ': dados_filtrados.iloc[:, 3],
                'Data': dados_filtrados.iloc[:, 0],
                'Cidade a recolher': dados_filtrados.iloc[:, 2],
                'Faturamento': dados_filtrados.iloc[:, 10]
            })
            
            # Filtrar valores nulos de faturamento para o cálculo do total
            resultado_com_faturamento = resultado[pd.notna(resultado['Faturamento'])]
            
            # Calcular o valor total de faturamento (apenas valores não nulos)
            valor_total = resultado_com_faturamento['Faturamento'].sum()
            valor_total_formatado = self.formatar_valor_brl(valor_total)
            
            # Formatar datas para exibição
            data_inicial_formatada = self.formatar_data_abreviada(data_inicial)
            data_final_formatada = self.formatar_data_abreviada(data_final)
            
            # Obter data atual formatada
            data_atual_formatada = self.obter_data_atual_formatada()
            
            # Perguntar onde salvar o documento Word
            caminho_salvar = filedialog.asksaveasfilename(
                title="Salvar documento Word",
                defaultextension=".docx",
                filetypes=[("Arquivos Word", "*.docx")]
            )
            
            if not caminho_salvar:
                self.status_var.set("Operação cancelada pelo usuário")
                return
            
            # Carregar o modelo Word
            try:
                doc = Document(self.caminho_modelo_word.get())
            except Exception as e:
                messagebox.showerror("Erro", f"Não foi possível abrir o modelo Word: {str(e)}")
                self.status_var.set("Erro ao abrir modelo Word")
                return
            
            # Substituir marcadores
            self.substituir_texto_word(doc, "<<CNPJ>>", cnpj_filtrado)
            self.substituir_texto_word(doc, "<<VALOR_TOTAL>>", valor_total_formatado.replace('R$', '').strip())
            self.substituir_texto_word(doc, "<<DATA_INICIAL>>", data_inicial_formatada)
            self.substituir_texto_word(doc, "<<DATA_FINAL>>", data_final_formatada)
            self.substituir_texto_word(doc, "<<DATA_CORRENTE>>", data_atual_formatada)
            
            # Criar tabela com os dados (a função já filtra valores nulos internamente)
            self.criar_tabela_resultados(doc, resultado)
            
            # Salvar documento Word
            doc.save(caminho_salvar)
            
            # Converter para PDF
            try:
                self.status_var.set("Convertendo para PDF...")
                self.root.update_idletasks()
                
                caminho_pdf = self.converter_para_pdf(caminho_salvar)
                self.status_var.set("Documentos salvos com sucesso!")
                messagebox.showinfo("Sucesso", f"Documentos processados e salvos:\n- Word: {caminho_salvar}\n- PDF: {caminho_pdf}")
                
            except Exception as e:
                self.status_var.set(f"Documento Word salvo. Erro ao gerar PDF: {str(e)}")
                messagebox.showinfo("Sucesso parcial", f"Documento Word salvo em:\n{caminho_salvar}\n\nErro ao gerar PDF: {str(e)}")
                
        except Exception as e:
            messagebox.showerror("Erro", f"Ocorreu um erro ao processar os dados:\n{str(e)}")
            self.status_var.set("Erro no processamento")

def main():
    # Verificar bibliotecas necessárias
    bibliotecas_faltando = []
    
    try:
        import docx
    except ImportError:
        bibliotecas_faltando.append("python-docx")
    
    # Verificar bibliotecas para PDF
    pdf_disponivel = False
    try:
        import win32com.client
        pdf_disponivel = True
    except ImportError:
        try:
            import docx2pdf
            pdf_disponivel = True
        except ImportError:
            try:
                import pypandoc
                pdf_disponivel = True
            except ImportError:
                pass
    
    if bibliotecas_faltando:
        messagebox.showerror("Erro de Importação", 
                            f"Bibliotecas não encontradas: {', '.join(bibliotecas_faltando)}\n" +
                            "Instale usando: pip install " + " ".join(bibliotecas_faltando))
        return
    
    if not pdf_disponivel:
        messagebox.showwarning("Aviso", 
                              "Nenhuma biblioteca de conversão PDF encontrada.\n" +
                              "Instale uma das seguintes opções:\n" +
                              "- pip install pywin32 (recomendado para Windows)\n" +
                              "- pip install docx2pdf\n" +
                              "- pip install pypandoc\n\n" +
                              "O programa funcionará, mas não gerará PDF.")
    
    root = Window(themename="superhero")
    app = FaturamentoPromilApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()